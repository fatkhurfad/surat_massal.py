import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import zipfile

st.set_page_config(page_title="Generator Surat Massal", layout="centered")
st.title("📄 Generator Surat Massal TIM PMT")
st.markdown("""
    <style>
    html, body, .stApp {
        background-color: #f7f9fa;
        font-family: 'Helvetica Neue', 'Segoe UI', sans-serif;
        color: #212121;
    }

    .stButton > button, .stDownloadButton > button {
        font-size: 16px;
        font-weight: 600;
        padding: 10px 24px;
        border-radius: 8px;
        border: none;
        transition: background-color 0.3s ease;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }

    .stButton > button {
        background-color: #03AC0E; /* Tokopedia green */
        color: white;
    }

    .stButton > button:hover {
        background-color: #02960C;
    }

    .stDownloadButton > button {
        background-color: #028A0D;
        color: white;
    }

    .stDownloadButton > button:hover {
        background-color: #026D0A;
    }

    .stFileUploader label {
        color: #424242;
        font-weight: 600;
    }

    .css-10trblm {
        font-size: 24px;
        font-weight: 700;
        color: #03AC0E;
    }

    .stTextInput > div > input {
        border-radius: 6px;
        padding: 0.5em;
    }

    .stAlert {
        background-color: #e8f5e9;
        border-left: 5px solid #03AC0E;
    }
    </style>
""", unsafe_allow_html=True)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")

    # Styling hyperlink (Arial, 12pt, blue, underline)
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")  # 12pt = 24 half-points
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")  # blue
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

uploaded_template = st.file_uploader("Upload Template Word (.docx)", type="docx")
uploaded_excel = st.file_uploader("Upload Data Excel (.xlsx)", type="xlsx")

if uploaded_template and uploaded_excel:
    df = pd.read_excel(uploaded_excel)

    if not {'nama_penyelenggara', 'short_link'}.issubset(df.columns):
        st.error("Excel harus memiliki kolom: 'nama_penyelenggara' dan 'short_link'")
    else:
        if st.button("🔄 Generate Surat"):
            output_zip = BytesIO()
            with zipfile.ZipFile(output_zip, "w") as zf:
                for _, row in df.iterrows():
                    doc = Document(uploaded_template)

                    for p in doc.paragraphs:
                        for run in p.runs:
                            if "{{nama_penyelenggara}}" in run.text:
                                run.text = run.text.replace("{{nama_penyelenggara}}", row["nama_penyelenggara"])

                    for p in doc.paragraphs:
                        if "{{short_link}}" in p.text:
                            parts = p.text.split("{{short_link}}")
                            p.clear()
                            if parts[0]: p.add_run(parts[0])
                            add_hyperlink(p, row["short_link"], row["short_link"])
                            if len(parts) > 1: p.add_run(parts[1])

                    for p in doc.paragraphs:
                        for run in p.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(12)

                    buffer = BytesIO()
                    filename = f"{row['nama_penyelenggara'].replace('/', '-')}.docx"
                    doc.save(buffer)
                    zf.writestr(filename, buffer.getvalue())

            st.success("✅ Surat berhasil dibuat!")
            st.download_button(
                label="📥 Download ZIP Hasil",
                data=output_zip.getvalue(),
                file_name="surat_massal_output.zip",
                mime="application/zip"
            )
